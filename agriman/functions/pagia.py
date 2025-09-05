from docx import Document
from io import BytesIO
from fastapi.responses import StreamingResponse

import pandas as pd
from datetime import date
from docxtpl import DocxTemplate
from agriman.database import get_engine

def get_pagia(customer_id, period_id):

  engine = get_engine()

  query = f"""
  SELECT
    applications.afm,
    applications.firstname,
    applications.lastname,
    applications.fathername,
    applications.idno,
    applications.phone,
    applications.mobil,
    applications.year,
    applications.iban,
    applications.tk,
    applications.lkkoi_b2_description,
    banks.code AS bankscode
  FROM applications
  JOIN customers ON customers.afm = applications.afm
  JOIN banks ON banks.id = applications.bank_id
  WHERE applications.period_id = {period_id} AND
        customers.id = {customer_id}
  """
  df=pd.read_sql(query, con=engine)
  
  if df.loc[0,'bankscode'] == '017': ##ΤΡΑΠΕΖΑ ΠΕΙΡΑΙΩΣ Α.Ε. (017)
    doc = DocxTemplate("./agriman/templates/pagiaTP.docx")
  elif df.loc[0,'bankscode'] == '011': ##ΕΘΝΙΚΗ ΤΡΑΠΕΖΑ ΤΗΣ ΕΛΛΑΔΟΣ Α.Ε. (011)
    doc = DocxTemplate("./agriman/templates/pagiaNBG.docx")
  else:
    doc = Document()
    doc.add_heading('ΠΡΟΣΟΧΗ', level=0)
    doc.add_paragraph('Ο πελάτης δεν διαθέτει τραπεζικό λογαριασμό είτε στην ΤΡΑΠΕΖΑ ΠΕΙΡΑΙΩΣ Α.Ε. είτε ΕΘΝΙΚΗ ΤΡΑΠΕΖΑ ΤΗΣ ΕΛΛΑΔΟΣ Α.Ε. ή δεν έχει δηλώσει τραπεζικό λογαριασμό')
  
  context ={
    'firstname' : df.loc[0,'firstname'],
    'lastname' : df.loc[0,'lastname'],
    'fathername' : df.loc[0,'fathername'],
    'afm' : df.loc[0,'afm'],
    'idno' : df.loc[0,'idno'],
    'tk' : df.loc[0,'tk'],
    'lkkoi_b2_description' : df.loc[0,'lkkoi_b2_description'],
    'phone' : df.loc[0,'phone'],
    'mobil' : df.loc[0,'mobil'],
    'iban' : df.loc[0,'iban'],
    'bankscode' : df.loc[0,'bankscode'],
    'year' : df.loc[0,'year'],
    'date' : date.today().strftime("%d/%m/%Y")
  }
  doc.render(context)

  buffer = BytesIO()
  doc.save(buffer)
  buffer.seek(0)

  # return as a downloadable file
  return StreamingResponse(
    buffer,
    media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    headers={
       "Content-Disposition": f"attachment; filename=pagia_{customer_id}.docx"
    }
  )
