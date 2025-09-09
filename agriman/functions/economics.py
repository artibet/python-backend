from docx import Document
from io import BytesIO
from fastapi.responses import StreamingResponse
from babel.numbers import format_currency

import pandas as pd
from datetime import date
from docxtpl import DocxTemplate
from agriman.database import get_engine

engine = get_engine()

def fun_details(id_key):
	#### SQL ερώτημα
	query = f"""
	SELECT
	    applications.afm,
	    applications.firstname,
	    applications.lastname,
	    applications.year,
	    applications.esap,
	    applications.book_number,
	    applications.dikaiomata_total,
	    banks.descr AS bank,
	    applications.iban
	FROM applications 
	JOIN banks ON applications.bank_id = banks.id
	WHERE applications.id = '{id_key}' 
	"""
	return(pd.read_sql(query, con=engine))


def fun_parcels(id_key):
    query = f"""
    SELECT
        applications.afm,
        applications.year,
        parcels.application_id,
        COUNT(*) AS num_parcels
    FROM parcels 
    JOIN applications ON applications.id = parcels.application_id
    WHERE applications.id = '{id_key}' AND parcels.is_pasture = 0
    GROUP BY parcels.application_id
    ORDER BY applications.afm, parcels.application_id
    """
    df = pd.read_sql(query, con=engine)
    if df.empty == True:
        num_parcels = 0
    else:
        num_parcels=df.loc[0,'num_parcels']
    return(num_parcels)

def fun_stables(id_key):
    query = f"""
    SELECT
        applications.afm,
        applications.year,
        stables.application_id,
        COUNT(*) AS num_stables
    FROM stables 
    JOIN applications ON applications.id = stables.application_id
    WHERE applications.id = '{id_key}'
    GROUP BY stables.application_id
    ORDER BY applications.afm, stables.application_id
    """
    df = pd.read_sql(query, con=engine)
    if df.empty == True:
        num_stables = 0
    else:
        num_stables=df.loc[0,'num_stables']
    return(num_stables)

def fun_equals(id_key):
    query = f"""
    SELECT
        applications.afm,
        applications.year,
        application_measures.application_id,
        COUNT(*) AS num_equals
    FROM application_measures 
    JOIN applications ON applications.id = application_measures.application_id
    WHERE applications.id = '{id_key}' AND application_measures.code IN ('13.1','13.2','13.3')
    GROUP BY application_measures.application_id
    ORDER BY applications.afm, application_measures.application_id
    """
    df = pd.read_sql(query, con=engine)
    if df.empty == True:
        num_equals = 0
    else:
        num_equals=df.loc[0,'num_equals']
    return(num_equals)

def fun_equals_parcel(id_key):
    query = f"""
    SELECT
        applications.afm,
        applications.year,
        measures.code AS code,
        SUM(parcels.area) AS area_per_measure,
        COUNT(*) AS num_measures
    FROM measures
    JOIN parcels ON parcels.id = measures.parcel_id
    JOIN applications ON applications.id = parcels.application_id
    WHERE applications.id = '{id_key}' AND measures.code IN ('13.1','13.2','13.3')
    GROUP BY measures.code
    ORDER BY applications.afm, measures.code
    """
    df = pd.read_sql(query, con=engine)
##    if df.empty == True:
##        equals = 0
##    else:
##        equals=df
    return(df)


def fun_equals_stable(id_key):
    query = f"""
    SELECT
        applications.afm,
        applications.year,
        stable_measures.code AS code,
        COUNT(*) AS num_stables
    FROM stable_measures
    JOIN stables ON stables.id = stable_measures.stable_id
    JOIN applications ON applications.id = stables.application_id
    WHERE applications.id = '{id_key}' AND stable_measures.code IN ('13.1','13.2','13.3')
    GROUP BY stable_measures.code
    ORDER BY applications.afm, stable_measures.code
    """
    df = pd.read_sql(query, con=engine)
##    if df.empty == True:
##        equals = 0
##    else:
##        equals=df
    return(df)

def fun_ecos(id_key):
    query = f"""
    SELECT
        applications.afm,
        applications.year,
        ecoschemes.code AS code,
        SUM(parcels.area) AS area_per_eco,
        COUNT(DISTINCT parcels.id) AS num_parcel
    FROM parcels
    JOIN applications ON applications.id = parcels.application_id
    JOIN ecoschemes ON ecoschemes.parcel_id = parcels.id
    WHERE applications.id = '{id_key}' AND parcels.is_seed_cultivation = 0
    GROUP BY applications.afm, ecoschemes.code
    ORDER BY ecoschemes.code
    """
    return(pd.read_sql(query, con=engine))

def fun_support(id_key):
    query = f"""
    SELECT
        applications.afm,
        applications.year,
        support_schemes.code AS code,
        supports.descr AS descr,
        SUM(parcel_cultivations.area) AS area_per_support,
        COUNT(DISTINCT parcel_cultivation_id) AS num_cultivations
    FROM support_schemes
    JOIN supports ON supports.code = support_schemes.code
    JOIN parcel_cultivations ON parcel_cultivations.id = support_schemes.parcel_cultivation_id
    JOIN parcels ON parcels.id = parcel_cultivations.parcel_id
    JOIN applications ON applications.id = parcels.application_id
    WHERE applications.id = '{id_key}' AND parcels.is_seed_cultivation = 0
    GROUP BY applications.afm, support_schemes.code, supports.descr
    ORDER BY support_schemes.code
    """
    return(pd.read_sql(query, con=engine))

def fun_measures(id_key):
    query = f"""
    SELECT
        applications.afm,
        applications.year,
        measures.code AS code,
        SUM(parcels.area) AS area_per_measure,
        COUNT(*) AS num_measures
    FROM measures
    JOIN parcels ON parcels.id = measures.parcel_id
    JOIN applications ON applications.id = parcels.application_id
    WHERE applications.id = '{id_key}' AND measures.code NOT IN ('13.1','13.2','13.3')
    GROUP BY measures.code
    ORDER BY applications.afm, measures.code
    """
    return(pd.read_sql(query, con=engine))

def fun_elga_cult(id_key):
    query = f"""
    SELECT
        applications.afm,
        applications.year,
        SUM(parcel_cultivations.cost) AS sum_cultivations,
        COUNT(DISTINCT parcel_cultivations.id) AS num_cultivations
    FROM parcel_cultivations
    JOIN parcels ON parcels.id = parcel_cultivations.parcel_id
    JOIN applications ON applications.id = parcels.application_id
    WHERE applications.id = '{id_key}' AND parcels.is_seed_cultivation = 0
    """
    df = pd.read_sql(query, con=engine)
    if pd.isna(df.loc[0, 'sum_cultivations']):
        value=0
    else:
        value=df.loc[0,'sum_cultivations']
    return(value)

def fun_elga_catl(id_key):
    query = f"""
    SELECT
        applications.afm,
        applications.year,
        SUM(cattles.cost) AS sum_cattles,
        COUNT(DISTINCT cattles.id) AS num_cattles
    FROM cattles
    JOIN stables ON stables.id = cattles.stable_id
    JOIN applications ON applications.id = stables.application_id
    WHERE applications.id = '{id_key}' 
    """
    df = pd.read_sql(query, con=engine)
    if pd.isna(df.loc[0, 'sum_cattles']):
        value=0
    else:
        value=df.loc[0,'sum_cattles']
    return(value)
    
def fun_elga_anim(id_key):
    query = f"""
    SELECT
        applications.afm,
        applications.year,
        SUM(animals.cost) AS sum_animals,
        COUNT(DISTINCT animals.id) AS num_animals
    FROM animals
    JOIN stables ON stables.id = animals.stable_id
    JOIN applications ON applications.id = stables.application_id
    WHERE applications.id = '{id_key}' 
    """
    df = pd.read_sql(query, con=engine)
    if pd.isna(df.loc[0, 'sum_animals']):
        value=0
    else:
        value=df.loc[0,'sum_animals']
    return(value)

def find_ecos_cost(value, year_key):
    query = f"""
    SELECT
        ecoscheme_credits.credit AS cost,
        periods.year AS year
    FROM ecoscheme_credits
    JOIN periods ON ecoscheme_credits.period_id=periods.id
    WHERE periods.year = '{year_key}' AND ecoscheme_credits.code = '{value}'
    """
    df = pd.read_sql(query, con=engine)
    return(df.loc[0,'cost'])

def find_equal_cost(value, year_key, area):
    query = f"""
    SELECT
        equalizer_credit_scales.credit AS cost,
        equalizer_credit_scales.scale AS aa,
        periods.year AS year
    FROM equalizer_credit_scales
    JOIN equalizer_credits ON equalizer_credit_scales.equalizer_credit_id=equalizer_credits.id
    JOIN periods ON periods.id = equalizer_credits.period_id
    WHERE periods.year = '{year_key}' AND equalizer_credits.code = '{value}'
    """
    df = pd.read_sql(query, con=engine)
    if area <= df.loc[0,'aa']:
        res=area*df.loc[0,'cost']
    elif area<= df.loc[0,'aa']+df.loc[1,'aa']:
        res=df.loc[0,'aa']*df.loc[0,'cost'] + (area-df.loc[0,'aa'])*df.loc[1,'cost']
    elif area<= df.loc[0,'aa']+df.loc[1,'aa']+df.loc[2,'aa']:
        res=df.loc[0,'aa']*df.loc[0,'cost'] + df.loc[1,'aa']*df.loc[1,'cost'] + (area-df.loc[0,'aa']-df.loc[1,'aa'])*df.loc[2,'cost']
    else:
        res=df.loc[0,'aa']*df.loc[0,'cost'] + df.loc[1,'aa']*df.loc[1,'cost'] + df.loc[2,'aa']*df.loc[2,'cost']
    return(res)

def find_parcel_cost(value, year_key):
    ## SQL ερώτημα
    query = f"""
    SELECT
        parcel_costs.from_parcels,
        parcel_costs.to_parcels,
        parcel_costs.cost,
        periods.year AS year
    FROM parcel_costs
    JOIN periods ON parcel_costs.period_id=periods.id
    WHERE periods.year = '{year_key}'
    ORDER BY parcel_costs.from_parcels
    """
    df = pd.read_sql(query, con=engine)
    cols=['from_parcels','to_parcels','cost']
    parcel_costs = df[cols].values.tolist()
    for i in range(len(parcel_costs)):
        if value >= parcel_costs[i][0] and value <= parcel_costs[i][1]:
            return(parcel_costs[i][2])
            break

def esap_cost(year_key):
    ## SQL ερώτημα
    query = f"""
    SELECT
        period_economics.esap
    FROM period_economics
    JOIN periods ON period_economics.period_id=periods.id
    WHERE periods.year= '{year_key}'
    """
    df = pd.read_sql(query, con=engine)
    if df.empty == True:
        esap_cost = 0
    else:
        esap_cost=df.loc[0,'esap']
    return(esap_cost)

def liters(id_key, status):
    if status==0:
        query = f"""
        SELECT
            SUM(liters)
        FROM v_diesel 
        JOIN applications ON applications.id =  v_diesel.application_id 
        WHERE applications.id = '{id_key}'
        """
    else:
        query = f"""
        SELECT
            variety_id, 
            variety_descr, 
            sum(liters)
        FROM v_diesel 
        JOIN applications ON applications.id =  v_diesel.application_id 
        WHERE applications.id = '{id_key}'
        GROUP BY variety_id, variety_descr
        """
    df = pd.read_sql(query, con=engine)
    return(df)







def get_economics(application_id):
  

  # doc = Document()
  # doc.add_heading('ΥΠΟ ΚΑΤΑΣΚΕΥΗ')
  # doc.add_paragraph('Το έγγραφο της οικονομικής ανάλυσης θα είναι διαθέσιμο σύντομα!')



  df1=fun_details(application_id)
  year_key=df1.loc[0,'year']
  esap_cost_val_f=esap_cost(year_key)
  esap_val = df1.loc[0, 'esap']
  if pd.isna(esap_val) or esap_val == "":
      esap_val = "---"
      esap_cost_val_f=0
  esap_cost_val=esap_cost_val_f/1.24
  p_c_f = 0
  if fun_parcels(afm, year_key)>0:
      p_c_f = find_parcel_cost(fun_parcels(application_id), year_key)
  p_c = p_c_f/1.24
  st_c_f = 0
  if fun_stables(application_id)>0:
      st_c_f = 10
  st_c = st_c_f/1.24
  eq_c_f = 0
  if fun_equals(application_id)>0:
      eq_c_f = 10
  eq_c = eq_c_f/1.24
  sum_d_f = p_c_f+st_c_f+eq_c_f
  sum_d = sum_d_f/1.24

  df2=fun_ecos(application_id)
  if df2 is None or df2.empty:
      ecoschemes = []   # κενό -> θα εμφανίσουμε "----" στο template
      ecos_total = []
  else:
      df2 = df2[['code', 'area_per_eco']].copy()
      # Στήλη με χρέωση
      df2['cost'] = df2['code'].apply(lambda c: find_ecos_cost(c, year_key))
      # Στήλη με γινόμενο
      df2['total_cost'] = df2['area_per_eco'] * df2['cost']
      ecos_total = df2['total_cost'].sum().round(2)
      # Στρογγυλοποίηση αν χρειάζεται
      df2['area_per_eco'] = df2['area_per_eco'].map(lambda x: f"{float(x):.2f}")
      df2['cost'] = df2['cost'].apply(eur_gr)
      df2['total_cost'] = df2['total_cost'].apply(eur_gr)
      ecoschemes = df2.to_dict(orient="records")
      ecos_total = eur_gr(ecos_total)
  # if df2.empty:
  #     ecoschemes="----"
  # else:
  #     ecoschemes= df2[['code','area_per_eco','num_parcel']]

  df3=fun_equals_parcel(application_id) 
  if df3 is None or df3.empty:
      equals_parcels = []   # κενό -> θα εμφανίσουμε "----" στο template
      equals_parcels_total = []
  else:
      df3 = df3[['code', 'area_per_measure']].copy()
      # Στήλη με χρέωση
      df3["cost"] = df3.apply(lambda row: find_equal_cost(row["code"], year_key, row["area_per_measure"]), axis=1)
      
      equals_parcels_total = df3['cost'].sum().round(2)
      # Στρογγυλοποίηση αν χρειάζεται
      df3['area_per_measure'] = df3['area_per_measure'].map(lambda x: f"{float(x):.2f}")
      df3['cost'] = df3['cost'].apply(eur_gr)
      equals_parcels = df3.to_dict(orient="records")
      equals_parcels_total = eur_gr(equals_parcels_total)


  # df3=fun_equals_parcel(afm, year_key)    
  # if df3.empty:
  #     equals_parcels="----"
  # else:
  #     equals_parcels = df3[['code','area_per_measure','num_measures']]
  #     print(afm)
  #     print(len(df3))
  #     print(df3.loc[0,'code'])
  #     print(find_equal_cost(df3.loc[0,'code'], year_key,df3.loc[0,'area_per_measure']))
  
  

  df4=fun_equals_stable(application_id)
  if df4.empty:
      equals_stables="----"
  else:
      equals_stables = df4[['code','num_stables']],

  # df5=fun_support(afm, year_key)
  # if df5.empty:
  #     support="----"
  # else:
  #     support = df5[['code','area_per_support','num_cultivations']]


  df5=fun_measures(application_id)
  if df5 is None or df5.empty:
      measures = []   # κενό -> θα εμφανίσουμε "----" στο template
      measures_total = []
  else:
      df5 = df5[['code', 'area_per_measure', 'num_measures']].copy()
      # Στήλη με χρέωση
      # df3["cost"] = df3.apply(lambda row: find_equal_cost(row["code"], year_key, row["area_per_measure"]), axis=1)
      
      # support_total = df3['cost'].sum().round(2)
      # Στρογγυλοποίηση αν χρειάζεται
      df5['area_per_measure'] = df5['area_per_measure'].map(lambda x: f"{float(x):.2f}")
      # df3['cost'] = df3['cost'].apply(eur_gr)
      measures = df5.to_dict(orient="records")
      # support_total = eur_gr(support_total)


  df7=fun_support(application_id)
  if df7 is None or df7.empty:
      support = []   # κενό -> θα εμφανίσουμε "----" στο template
      support_total = []
  else:
      df7 = df7[['code', 'descr', 'area_per_support']].copy()
      # Στήλη με χρέωση
      # df3["cost"] = df3.apply(lambda row: find_equal_cost(row["code"], year_key, row["area_per_measure"]), axis=1)
      
      # support_total = df3['cost'].sum().round(2)
      # Στρογγυλοποίηση αν χρειάζεται
      df7['area_per_support'] = df7['area_per_support'].map(lambda x: f"{float(x):.2f}")
      # df3['cost'] = df3['cost'].apply(eur_gr)
      support = df7.to_dict(orient="records")
      # support_total = eur_gr(support_total)
  
  df8=liters(application_id,0)

##
  data = {
      'firstname': df1.loc[0,'firstname'],
      'lastname': df1.loc[0,'lastname'],
      'afm': str(df1.loc[0,'afm']),
      'year': str(df1.loc[0,'year']),
      'book_number': str(df1.loc[0,'book_number']),
      'bank': df1.loc[0,'bank'],
      'iban': df1.loc[0,'iban'],
##
      'num_parcels': str(fun_parcels(afm, year_key)),
      'p_c_f': format_currency(p_c_f, 'EUR', locale='el_GR'),
      'p_c': format_currency(p_c, 'EUR', locale='el_GR'),
      'num_stables': str(fun_stables(afm, year_key)),
      'st_c_f': format_currency(st_c_f, 'EUR', locale='el_GR'),
      'st_c': format_currency(st_c, 'EUR', locale='el_GR'),
      'num_equal': str(fun_equals(afm, year_key)),
      'eq_c_f': format_currency(eq_c_f, 'EUR', locale='el_GR'),
      'eq_c': format_currency(eq_c, 'EUR', locale='el_GR'),
      'sum_d_f': format_currency(sum_d_f, 'EUR', locale='el_GR'),
      'sum_d': format_currency(sum_d, 'EUR', locale='el_GR'),
##        
      'esap': esap_val,
      'esap_cost_f': format_currency(esap_cost_val_f, 'EUR', locale='el_GR'),
      'esap_cost': format_currency(esap_cost_val, 'EUR', locale='el_GR'),
      'date': date.today().strftime("%d/%m/%Y"),


##
      'elga_cult': format_currency(fun_elga_cult(afm, year_key), 'EUR', locale='el_GR'),
      'elga_catl': format_currency(fun_elga_catl(afm, year_key), 'EUR', locale='el_GR'),
      'elga_anim': format_currency(fun_elga_anim(afm, year_key), 'EUR', locale='el_GR'),
      'elga_total': format_currency(fun_elga_cult(afm, year_key)+fun_elga_catl(afm, year_key)+fun_elga_anim(afm, year_key), 'EUR', locale='el_GR'),
##
      'dikaiomata_total': format_currency(df1.loc[0,'dikaiomata_total'], 'EUR', locale='el_GR'),
##
      'ecoschemes': ecoschemes,
      'ecos_total': ecos_total,
      'equals_parcels': equals_parcels,
      'equals_parcels_total': equals_parcels_total,
      'equals_stables': equals_stables,
      'measures': measures,
      'support': support,

      'liters': df8
  }

  doc = DocxTemplate("./agriman/templates/economics.docx")
  doc.render(data)
  # Save in memory
  buffer = BytesIO()
  doc.save(buffer)
  buffer.seek(0)

  # return as a downloadable file
  return StreamingResponse(
    buffer,
    media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    headers={
       "Content-Disposition": f"attachment; filename=economics_{application_id}.docx"
    }
  )
